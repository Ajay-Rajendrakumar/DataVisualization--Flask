from docx import Document
from docx.shared import Inches
from datetime import datetime
from zipfile import ZipFile
import glob
import glob, os

document = Document()	
dateTimeObj = datetime.now()
d = dateTimeObj.strftime("%b%d%Y")
cwd = os.getcwd() + '/static/Zips/'
zipObj = ZipFile(cwd+'Report-'+d+'.zip', 'w')
#if request.method == 'POST':
#	val = request.form['object']
imgfiles = []
cwd = os.getcwd() + '/static/Reports'

for file in os.listdir(cwd):
	imgfiles.append(file)
for i in imgfiles:
	p = document.add_paragraph()
	r = p.add_run()
	r.add_text(i)
	r.add_picture('static/Reports/'+i)
	zipObj.write('static/Reports/'+i)
	os.remove('static/Reports/'+i) 
cwd = os.getcwd() + '/static/Report-Docxs/'
document.save(cwd+val+'.docx')	
zipObj.close()