from flask import Flask, render_template, request
import matplotlib.pyplot as plt; plt.rcdefaults()
import plotly.graph_objects as go
import sqlite3 as sql
import pandas as pd
import sqlite3  
import numpy as np
import os
import random
from docx import Document
from docx.shared import Inches
import glob
from zipfile import ZipFile
from datetime import datetime

warn="no"
error="no"
Table_Name = "Parts_Info"
global count
count=0


app = Flask(__name__)

def Extract(value):
	global Table_Name

	con = sqlite3.connect("HW_Parts.db")
	cur = con.cursor()		
	sql = "SELECT DISTINCT " +value+", count("+value+") FROM "+Table_Name+" GROUP BY "+value+";"
	x=[]
	y=[]
	for row,n in cur.execute(sql):
		x.append(row)
		y.append(n)
	return x,y
def help():
	global Table_Name
	con = sqlite3.connect("HW_Parts.db")
	cur= con.cursor()
	cursor = con.execute('select * from '+Table_Name)
	names = list(map(lambda x: x[0], cursor.description))
	table=[]
	for i in cur.execute("SELECT name FROM sqlite_master WHERE type='table';"):
		table.append(i[0])
	return names,table

@app.route('/')
def home():
	return render_template('home.html')

@app.route('/dash')
def dash():
		cols,table=help()
		return render_template('Dashboard.html',cols=cols,tbname=table)

@app.route('/bar', methods=['GET', 'POST'])	
def bar():
	global count
	global warn
	global error
	global Table_Name
	fl='static/Reports/Barchart'+str(count)+'.png'
	count+=1
	if os.path.exists(fl):
		os.remove(fl)
		
	if request.method == 'POST':
		val = request.form['object']
	cols,table=help()
	x,y=Extract(val)
	y_pos = np.arange(len(x))
	plt.barh(y_pos, y, align='center', alpha=0.5,color=(0.2, 0.4, 0.6, 0.6))
	plt.yticks(y_pos, x)
	plt.xlabel(val)
	plt.title("Bar Chart")
	for index, value in enumerate(y):
		plt.text(value, index, str(value))
	total= len(x)
	if(total>10):
		warn="o"
	plt.savefig(fl,bbox_inches = 'tight')
	plt.clf()
	return render_template('Dashboard.html',l=x,g=y,f=fl,res="Barchart",head=val,warn=warn,cols=cols,tbname=table)

@app.route('/pie', methods=['GET', 'POST'])
def pie():
	global Table_Name
	global warn
	global error
	global count
	fl='static/Reports/Piechart'+str(count)+'.png'
	count+=1
	if os.path.exists(fl):
		os.remove(fl)
	if request.method == 'POST':
		val = request.form['object']
	cols,table=help()
	x,y=Extract(val)
	number_of_colors = len(x)
	total= len(x)
	if(total>10):
		warn="o"
	color = ["#"+''.join([random.choice('0123456789ABCDEF') for j in range(6)])for i in range(number_of_colors)]
	plt.pie(y, labels=x, colors=color,autopct='%1.1f%%', shadow=True, startangle=140)
	plt.axis('equal')
	plt.savefig(fl,bbox_inches = 'tight')
	plt.clf()
	return render_template('Dashboard.html',l=x,g=y,f=fl,res="Piechart",head=val,warn=warn,cols=cols,error=error,tbname=table)

@app.route('/sca', methods=['GET', 'POST'])
def sca():
	global Table_Name
	global warn
	global error
	global count
	fl='static/Reports/Scatter-plot-chart'+str(count)+'.png'
	count+=1
	if os.path.exists(fl):
		os.remove(fl)
	if request.method == 'POST':
		val = request.form['object']
	cols,table=help()
	if val not in cols:
		return render_template('Dashboard.html',error="yes",head=val,cols=cols,tbname=table)
	x,y=Extract(val)
	total= len(x)
	if(total>10):
		warn="o"
	plt.scatter(x, y) 
	plt.savefig(fl,bbox_inches = 'tight')
	plt.clf()
	return render_template('Dashboard.html',l=x,g=y,f=fl,res="Scatter-plot-chart",head=val,warn=warn,cols=cols,error=error,tbname=table)

@app.route('/line', methods=['GET', 'POST'])
def line():
	global warn
	global error
	global Table_Name
	global count
	fl='static/Reports/Line-plot-chart'+str(count)+'.png'
	count+=1
	if os.path.exists(fl):
		os.remove(fl)
	cols,table=help()
	if request.method == 'POST':
		val = request.form['object']
	cols,table=help()
	x,y=Extract(val)
	total= len(x)
	if(total>10):
		warn="o"
	plt.plot(x,y) 
	plt.savefig(fl,bbox_inches = 'tight')
	plt.clf()
	return render_template('Dashboard.html',l=x,g=y,f=fl,res="Line-plot-chart",head=val,warn=warn,cols=cols,error=error,tbname=table)

@app.route('/view', methods=['GET', 'POST'])
def view():
	global Table_Name
	con = sqlite3.connect("HW_Parts.db")
	cur=con.execute("SELECT * FROM "+Table_Name)
	data = cur.fetchall()
	names = list(map(lambda x: x[0], cur.description))
	cols,table=help()
	return render_template('view.html', data=data,clname=names,cols=names,tbname=table)


@app.route('/query', methods=['GET', 'POST'])
def query():
	con = sqlite3.connect("HW_Parts.db")
	if request.method == 'POST':
		val = request.form['object']

	cur=con.execute(val)
	data = cur.fetchall()
	cols,table=help()
	names = list(map(lambda x: x[0], cur.description))	
	return render_template('view.html', data=data,clname=names,cols=cols,tbname=table)

@app.route('/report', methods=['GET', 'POST'])
def report():
	document = Document()	
	dateTimeObj = datetime.now()
	d = dateTimeObj.strftime("%b%d%Y")
	cwd = os.getcwd() + '/static/Zips/'
	zipObj = ZipFile(cwd+'Report-'+d+'.zip', 'w')
	if request.method == 'POST':
		val = request.form['object']
	imgfiles = []
	cwd = os.getcwd() + '/static/Reports'

	for file in os.listdir(cwd):
		imgfiles.append(file)
	for i in imgfiles:
		p = document.add_paragraph()
		r = p.add_run()
		r.add_text(i)
		r.add_picture('static/Reports/'+i)
		zipObj.write('static/Reports/'+i)
		os.remove('static/Reports/'+i) 
	cwd = os.getcwd() + '/static/Report-Docxs/'
	document.save(cwd+val+'.docx')	
	zipObj.close()
	return render_template('home.html')

if __name__ == '__main__':
	app.run(debug = True)