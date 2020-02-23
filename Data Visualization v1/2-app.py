from flask import Flask, render_template, request
import matplotlib.pyplot as plt; plt.rcdefaults()
from docx.shared import Inches
from datetime import datetime
from zipfile import ZipFile
from docx import Document
import pandas as pd
import numpy as np
import sqlite3 
import random
import glob
import os

TABLE_NAME= 'PartsInfo'
TABLE_NAME1='Employee'
DB = 'HW_Parts.db'
WARN='no'
COUNT=0


app = Flask(__name__)

def Extract(value):
	con = sqlite3.connect(DB)
	cur = con.cursor()
	x=[]
	y=[]
	
	if(value=="CreatedEmpId" or value=="ApprovedEmpId"):
		sql ="SELECT DISTINCT  Employee.EmpName , count(PartsInfo."+value +") FROM Employee INNER JOIN PartsInfo ON PartsInfo."+value+"=Employee.EmpId GROUP BY PartsInfo."+ value +";"
	else:
		sql = "SELECT DISTINCT " +value+", count("+value+") FROM "+ TABLE_NAME +" GROUP BY "+value+";"
	for name,count in cur.execute(sql):
		x.append(name)
		y.append(count)		
		
	return x,y

def help():
	con = sqlite3.connect(DB)
	cur= con.cursor()
	cursor = con.execute('select * from '+ TABLE_NAME)
	Colnames = list(map(lambda x: x[0], cursor.description))
	table=[]
	for i in cur.execute("SELECT name FROM sqlite_master WHERE type='table';"):
		table.append(i[0])
	return Colnames,table

@app.route('/')
def home():
	return render_template('home.html')

@app.route('/dash')
def dash():
		cols,tables=help()
		return render_template('Dashboard.html',cols=cols,tbname=tables)

@app.route('/bar', methods=['GET', 'POST'])	
def bar():
	global COUNT
	if request.method == 'POST':
		val = request.form['object']
	fl='static/Reports/Barchart-'+str(COUNT)
	fl=fl+"-Cat-"+val+'.png'
	COUNT+=1
	if os.path.exists(fl):
		os.remove(fl)

	cols,tables=help()
	
	x,y=Extract(val)
	y_pos = np.arange(len(x))
	plt.barh(y_pos, y, align='center', alpha=0.5,color=(0.2, 0.4, 0.6, 0.6))
	plt.yticks(y_pos, x)
	plt.xlabel(val)
	plt.title("Bar Chart")
	for index, value in enumerate(y):
		plt.text(value, index, str(value))
	total= len(x)
	
	if(total>10):
		WARN="o"
	else:
		WARN="no"
	plt.savefig(fl,bbox_inches = 'tight')
	plt.clf()
	return render_template('Dashboard.html',l=x,g=y,f=fl,res="Barchart",head=val,warn=WARN,cols=cols,tbname=tables)



@app.route('/pie', methods=['GET', 'POST'])
def pie():
	global COUNT
	if request.method == 'POST':
		val = request.form['object']
	fl='static/Reports/Piechart-'+str(COUNT)
	fl=fl+"-Cat-"+val+'.png'
	COUNT+=1
	if os.path.exists(fl):
		os.remove(fl)
	cols,tables=help()
	
	x,y=Extract(val)
	s=[]
	for i in range(len(x)):
		s.append(x[i]+","+str(y[i]))

	number_of_colors = len(x)
	total= len(x)
	if(total>10):
		WARN="o"
	else:
		WARN="no"
	color = ["#"+''.join([random.choice('0123456789ABCDEF') for j in range(6)])for i in range(number_of_colors)]
	plt.pie(y, labels=s, colors=color,autopct='%1.1f%%', shadow=True, startangle=140)
	plt.axis('equal')
	plt.savefig(fl,bbox_inches = 'tight')
	plt.clf()
	return render_template('Dashboard.html',l=x,g=y,f=fl,res="Piechart",head=val,warn=WARN,cols=cols,tbname=tables)


@app.route('/sca', methods=['GET', 'POST'])
def sca():
	global COUNT
	if request.method == 'POST':
		val = request.form['object']
	fl='static/Reports/Scatter-plot-chart-'+str(COUNT)
	fl=fl+"-Cat-"+val+'.png'
	COUNT+=1
	if os.path.exists(fl):
		os.remove(fl)
	cols,tables=help()
	
	x,y=Extract(val)
	total= len(x)
	if(total>10):
		WARN="o"
	else:
		WARN="no"
	plt.scatter(x, y) 
	plt.savefig(fl,bbox_inches = 'tight')
	plt.clf()
	return render_template('Dashboard.html',l=x,g=y,f=fl,res="Scatter-plot-chart",head=val,warn=WARN,cols=cols,tbname=tables)


@app.route('/line', methods=['GET', 'POST'])
def line():
	global COUNT
	if request.method == 'POST':
		val = request.form['object']
	fl='static/Reports/Line-plot-chart-'+str(COUNT)
	fl=fl+"-Cat-"+val+'.png'
	COUNT+=1
	if os.path.exists(fl):
		os.remove(fl)
	
	cols,tables=help()
	x,y=Extract(val)
	total= len(x)
	if(total>10):
		WARN="o"
	else:
		WARN="no"
	plt.plot(x,y) 
	plt.savefig(fl,bbox_inches = 'tight')
	plt.clf()
	return render_template('Dashboard.html',l=x,g=y,f=fl,res="Line-plot-chart",head=val,warn=WARN,cols=cols,tbname=tables)



@app.route('/view', methods=['GET', 'POST'])
def view():
	con = sqlite3.connect(DB)
	if request.method == 'POST':
		val = request.form['object']
	cur=con.execute("SELECT * FROM "+val)
	data = cur.fetchall()
	names = list(map(lambda x: x[0], cur.description))
	cols,tables=help()
	return render_template('view.html', data=data,clname=names,cols=names,tbname=tables)


@app.route('/query', methods=['GET', 'POST'])
def query():
	con = sqlite3.connect(DB)
	if request.method == 'POST':
		val = request.form['object']
	cur=con.execute(val)
	data = cur.fetchall()
	cols,table=help()
	names = list(map(lambda x: x[0], cur.description))	
	return render_template('view.html', data=data,clname=names,cols=cols,tbname=table)


@app.route('/report', methods=['GET', 'POST'])
def report():
	document = Document()	
	dateTimeObj = datetime.now()
	d = dateTimeObj.strftime("%b-%d-%Y-%H-%m")
	cwd = os.getcwd() + '/static/Zips/'
	zipObj = ZipFile(cwd+'Report-'+d+'.zip', 'w')
	if request.method == 'POST':
		val = request.form['object']
	imgfiles = []
	cwd = os.getcwd() + '/static/Reports'

	for file in os.listdir(cwd):
		imgfiles.append(file)
	if(len(imgfiles)==0):
		return render_template('home.html')
	else:	
		p = document.add_paragraph()
		r = p.add_run()
		r.add_text("Generated Reports are listed below:")
		for i in imgfiles:
			r.add_text(i[:-4])
			r.add_picture('static/Reports/'+i)
			zipObj.write('static/Reports/'+i)
			os.remove('static/Reports/'+i) 
		cwd = os.getcwd() + '/static/Report-Documents/'
		document.save(cwd+val+'.docx')	
		zipObj.close()
		return render_template('home.html',res="done")

if __name__ == '__main__':
	app.run(debug = True)




